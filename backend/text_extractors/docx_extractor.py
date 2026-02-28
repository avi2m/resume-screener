"""
DOCX Text Extraction Module
"""
import io
import logging

logger = logging.getLogger(__name__)


def extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX file content."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        
        parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        
        text = "\n".join(parts)
        logger.info(f"Extracted {len(text)} chars from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")
